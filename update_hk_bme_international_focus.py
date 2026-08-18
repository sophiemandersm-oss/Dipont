from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


DOCX = Path("hong_kong_biomedical_engineering_undergraduate_guide.docx")
LOGO = Path("dipont-logo 1.png")


PROGRAMMES = [
    {
        "uni": "HKU",
        "course": "BEng in Biomedical Engineering (JS6925)",
        "route": "Direct BME programme jointly offered by Engineering and HKUMed.",
        "alevel": "Competitive score: 2A*1A, excluding Chinese and English language subjects. Must include grade A or above in Mathematics, Further Mathematics or Pure Mathematics, and grade A or above in Biology, Chemistry or Physics.",
        "ib": "Competitive score: 36/45. Must include grade 6 or better in HL Mathematics: Analysis and Approaches or HL Mathematics: Applications and Interpretation, plus either grade 5 or better in HL Biology/Chemistry/Physics or grade 6 or better in SL Biology/Chemistry/Physics.",
        "ap": "HKU general lower boundary: grade 3 or above in three AP subjects, plus SAT 1380+ or grade 3 or above in two additional AP subjects. BME also requires high-school-level Mathematics and Biology, Chemistry or Physics, so applicants should present AP Calculus plus an AP laboratory science where possible.",
        "rate": "28.4%",
    },
    {
        "uni": "CUHK",
        "course": "BEng in Biomedical Engineering (JS4460)",
        "route": "Engineering programme offered with Medicine collaboration; focus areas include instrumentation, imaging/informatics/modelling, and molecular/cell/tissue engineering.",
        "alevel": "General minimum: 3 AL passes or 2 AL + 2 AS passes. BME preference: good grades in at least two of Physics, Chemistry, Biology or Mathematics.",
        "ib": "General minimum: IB Diploma with 30/45. BME preference: good grades in at least two of Physics, Chemistry, Biology or Mathematics.",
        "ap": "General minimum for US-pattern route: high-school diploma plus SAT 1190+ or ACT 24+, and grade 3 in two AP tests. BME preference: strong Mathematics and science APs, ideally two from Biology, Chemistry, Physics and Calculus.",
        "rate": "17.9%",
    },
    {
        "uni": "HKUST",
        "course": "BEng in Bioengineering via Chemical and Biological Engineering (JS5220); related BSc Biomedical and Health Sciences",
        "route": "Related route rather than a programme titled Biomedical Engineering. Bioengineering is the engineering route; Biomedical and Health Sciences is a science/health route.",
        "alevel": "Engineering/Bioengineering route: pass at least three AL subjects; 2025 intake mid-50% score range was AAA to 3A*. Subject requirement: senior-level Mathematics plus one senior-level subject from Physics, Chemistry, Biology or Computer Science. Physics or Chemistry is preferred for Chemical and Biological Engineering.",
        "ib": "General requirement: IB Diploma; 2025 intake mid-50% score range was 35-40 including bonus points. Subject requirement for Bioengineering: senior-level Mathematics plus one senior-level subject from Physics, Chemistry, Biology or Computer Science.",
        "ap": "General requirement: high-school graduation plus either SAT 1190+ or ACT 24+ with AP grade 3+ in two subjects, or AP grade 4+ in five subjects. Engineering/Bioengineering subject expectation: senior-level Mathematics plus a senior-level science; AP Calculus and AP Physics/Chemistry/Biology are the natural fit.",
        "rate": "10.0%",
    },
    {
        "uni": "PolyU",
        "course": "BSc (Hons) in Biomedical Engineering (JS3150)",
        "route": "Accredited BME programme with engineering common-year structure and BME/prosthetics-and-orthotics-related options.",
        "alevel": "International/non-local guideline: grade B or above in three AL subjects. Local Non-JUPAS page notes admitted A Level applicants typically attain at least 3B. Good results in science subjects such as Biology, Chemistry, Mathematics and/or Physics are preferred for JS3150.",
        "ib": "IB Diploma. Successful applicants typically achieved 32/45 or higher in recent years; predicted 30+ may be considered for interview. Science subjects are preferred for JS3150.",
        "ap": "US-pattern guideline: high-school diploma average 80%+ plus either SAT 1190+ or ACT 24+ and AP grade 3+ in two subjects, or AP grade 3+ in five subjects. Calculus AB/BC/Precalculus count as one AP subject; several AP subjects are excluded by PolyU. For BME, use AP Calculus plus Biology/Chemistry/Physics where possible.",
        "rate": "9.7%",
    },
    {
        "uni": "CityU",
        "course": "BEng Biomedical Engineering (JS1211)",
        "route": "Direct BME programme with features in medical technology, bioinstrumentation, cell and tissue engineering, and biomedical robotics.",
        "alevel": "General minimum: grade E or above in three GCE/International A Level subjects; two AS subjects may count as one AL subject. Programme expectation: science or engineering background; Physics and Mathematics backgrounds are preferred.",
        "ib": "General minimum for first-year entry: IB Diploma. For Advanced Standing I, CityU states a minimum diploma point score of 30/45. Programme expectation: science or engineering background; Physics and Mathematics preferred.",
        "ap": "CityU lists other secondary qualifications and country/test routes case-by-case on international admissions pages. For BME, applicants should satisfy CityU general entrance requirements and present a science/engineering background; AP Calculus plus AP Physics/Chemistry/Biology is the strongest fit.",
        "rate": "9.1%",
    },
]

SNAPSHOT_ROWS = [
    ["HKU", "BEng BME (JS6925)", "Direct BME; Engineering + HKUMed", "28.4%"],
    ["CUHK", "BEng BME (JS4460)", "Engineering BME; medicine collaboration", "17.9%"],
    ["HKUST", "Bioengineering/CBE (JS5220)", "Related engineering route", "10.0%"],
    ["PolyU", "BSc BME (JS3150)", "Accredited BME; applied engineering", "9.7%"],
    ["CityU", "BEng BME (JS1211)", "Direct BME; medical technology", "9.1%"],
]

TOC_ENTRIES = [
    ("1. Programme Snapshot and Admissions Rates", "1"),
    ("2. Course Summaries", "2"),
    ("3. Application Deadlines", "3"),
    ("4. Personal Statement Requirements", "4"),
    ("5. A Level Requirements", "5"),
    ("6. IB Requirements", "6"),
    ("7. AP / SAT Requirements", "7"),
    ("8. Planning Takeaways", "8"),
]

COURSE_SUMMARY_ROWS = [
    ["HKU", "BEng in Biomedical Engineering (JS6925)", "A direct biomedical engineering degree combining core engineering with medical and life-science context through HKU Engineering and HKUMed. Best suited to students who can show strong mathematics plus a physical or biological science, and who want a selective research-intensive BME route."],
    ["CUHK", "BEng in Biomedical Engineering (JS4460)", "An engineering-based BME route with explicit collaboration between Engineering and Medicine. The curriculum is a good fit for students interested in instrumentation, biomedical imaging, modelling, medical informatics, molecular engineering, cell engineering or tissue engineering."],
    ["HKUST", "BEng in Bioengineering via Chemical and Biological Engineering (JS5220)", "A related engineering pathway rather than a course titled Biomedical Engineering. It suits students interested in biological systems, chemical/biological engineering, biomolecular processes, and engineering applications in health, biotechnology and life sciences."],
    ["PolyU", "BSc (Hons) in Biomedical Engineering (JS3150)", "An accredited biomedical engineering programme with an applied, professionally oriented engineering feel. It is especially relevant for students interested in medical devices, rehabilitation engineering, prosthetics and orthotics, clinical technology and hands-on problem solving."],
    ["CityU", "BEng Biomedical Engineering (JS1211)", "A direct BME programme focused on medical technology, bioinstrumentation, cell and tissue engineering, biomedical robotics and healthcare engineering. It is a practical choice for students with a science or engineering background who want a focused BME route in Hong Kong."],
]

DEADLINE_ROWS = [
    ["HKU", "26 Nov 2025, 12:00 noon HKT", "21 Aug 2026, 12:00 noon HKT", "Applications after first round are reviewed on a rolling basis subject to places."],
    ["CUHK", "13 Nov 2025, 11:59 p.m. HKT", "8 Jan 2026 regular; 29 May 2026 extended", "One application covers later rounds; unsuccessful advance-offer applicants move into regular consideration."],
    ["HKUST", "20 Nov 2025", "8 Jan 2026 main; 30 Jun 2026 late", "Late applications are reviewed on a rolling basis after the main round."],
    ["PolyU", "19 Nov 2025 early for local Non-JUPAS", "5 Feb 2026 Non-JUPAS Year 1; 15 May 2026 international/other qualifications", "For JS3150, PolyU lists separate deadlines for Non-JUPAS Year 1 and International/Other Qualification routes."],
    ["CityU", "15 Nov 2025", "15 Jan 2026 main; late round from 16 Jan 2026", "Applications open 25 Sept 2025; interviews/offers run from November onward where applicable."],
]

PERSONAL_STATEMENT_ROWS = [
    ["HKU", "Required. One personal statement of no more than 1,000 words, regardless of programme choices, unless a special programme asks for a separate statement.", "For BME, focus on why engineering plus medicine/healthcare, evidence of maths/science readiness, and one or two technical or clinical-exposure examples."],
    ["CUHK", "Required for non-JUPAS/international applicants. Programmes other than Medicine require a personal statement up to two A4 pages in English.", "For CUHK BME, connect engineering thinking with medicine collaboration, and show readiness in at least two of Physics, Chemistry, Biology or Mathematics."],
    ["HKUST", "Required supporting document in the online application; HKUST lists a personal statement among documents to upload. No fixed length was stated on the reviewed admissions page.", "For Bioengineering/CBE, make the first-choice rationale explicit and link senior-level Mathematics plus science preparation to biological/chemical engineering applications."],
    ["PolyU", "Usually optional for most programmes. PolyU states recommendation letters and personal statements are optional unless specifically requested; no specific format is required.", "Use a concise statement anyway for JS3150: suitability for biomedical engineering, applied project evidence, and interest in BME/prosthetics/rehabilitation pathways."],
    ["CityU", "Listed as a document to prepare/upload for undergraduate application; the reviewed CityU pages did not publish a fixed word limit.", "For CityU BME, emphasise science/engineering background, medical-technology interest, practical projects, and readiness for interview or programme assessment if invited."],
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def style_table(table, font_size=9.2, cell_margin=100):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row_idx, row in enumerate(table.rows):
        set_cant_split(row)
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell, top=cell_margin, bottom=cell_margin)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(3)
                for run in paragraph.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(font_size)
            if row_idx == 0:
                set_cell_shading(cell, "E8EEF5")
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(31, 77, 120)
    set_repeat_table_header(table.rows[0])


def set_col_widths(table, widths):
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width


def add_table(doc, headers, rows, widths, font_size=9.2, cell_margin=100):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for idx, text in enumerate(headers):
        hdr[idx].text = text
    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            cells[idx].text = text
    set_col_widths(table, widths)
    style_table(table, font_size=font_size, cell_margin=cell_margin)
    return table


def add_note(doc, text):
    paragraph = doc.add_paragraph(text)
    paragraph.paragraph_format.left_indent = Inches(0.18)
    paragraph.paragraph_format.right_indent = Inches(0.18)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    for run in paragraph.runs:
        run.font.italic = True
        run.font.color.rgb = RGBColor(85, 85, 85)
    return paragraph


def add_page_break_if_needed(doc, needed=True):
    if needed:
        doc.add_page_break()


def add_toc_entry(doc, title, page):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.left_indent = Inches(0.15)
    paragraph.paragraph_format.right_indent = Inches(0.15)
    paragraph.paragraph_format.space_after = Pt(10)
    paragraph.paragraph_format.tab_stops.add_tab_stop(
        Inches(6.2), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS
    )
    run = paragraph.add_run(title)
    run.font.size = Pt(11)
    run.font.name = "Calibri"
    paragraph.add_run("\t")
    page_run = paragraph.add_run(page)
    page_run.font.size = Pt(11)
    page_run.font.name = "Calibri"


def add_numbered_heading(doc, number, title):
    heading = doc.add_heading(f"{number}. {title}", level=1)
    heading.paragraph_format.keep_with_next = True
    return heading


doc = Document()

section = doc.sections[0]
section.top_margin = Inches(0.72)
section.bottom_margin = Inches(0.72)
section.left_margin = Inches(0.78)
section.right_margin = Inches(0.78)

styles = doc.styles
styles["Normal"].font.name = "Calibri"
styles["Normal"].font.size = Pt(11)
styles["Normal"].paragraph_format.space_after = Pt(6)
styles["Normal"].paragraph_format.line_spacing = 1.15
styles["Title"].font.name = "Calibri"
styles["Title"].font.size = Pt(28)
styles["Title"].font.bold = True
styles["Title"].font.color.rgb = RGBColor(11, 37, 69)
styles["Subtitle"].font.name = "Calibri"
styles["Subtitle"].font.size = Pt(13)
styles["Subtitle"].font.color.rgb = RGBColor(31, 77, 120)
for style_name, size, color in [
    ("Heading 1", 16, RGBColor(46, 116, 181)),
    ("Heading 2", 13, RGBColor(46, 116, 181)),
    ("Heading 3", 12, RGBColor(31, 77, 120)),
]:
    style = styles[style_name]
    style.font.name = "Calibri"
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = color
    style.paragraph_format.space_before = Pt(6)
    style.paragraph_format.space_after = Pt(3)
    style.paragraph_format.keep_with_next = True

# Cover page
for _ in range(3):
    doc.add_paragraph()

logo_para = doc.add_paragraph()
logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
logo_para.paragraph_format.space_after = Pt(16)
logo_para.add_run().add_picture(str(LOGO), width=Inches(2.75))

brand = doc.add_paragraph()
brand.alignment = WD_ALIGN_PARAGRAPH.CENTER
brand_run = brand.add_run("Deep Intelligence")
brand_run.bold = True
brand_run.font.size = Pt(18)

title = doc.add_paragraph(style="Title")
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_before = Pt(16)
title.paragraph_format.space_after = Pt(18)
title.add_run("Hong Kong Biomedical Engineering\nUndergraduate Admissions Guide")

student = doc.add_paragraph()
student.alignment = WD_ALIGN_PARAGRAPH.CENTER
student_run = student.add_run("A Level, IB and AP Route Planning")
student_run.bold = True
student_run.font.size = Pt(16)

subtitle = doc.add_paragraph(style="Subtitle")
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.paragraph_format.space_before = Pt(16)
subtitle.add_run(
    "Top five Hong Kong universities | Biomedical engineering and related undergraduate courses"
)

summary = doc.add_paragraph()
summary.alignment = WD_ALIGN_PARAGRAPH.CENTER
summary.paragraph_format.left_indent = Inches(0.45)
summary.paragraph_format.right_indent = Inches(0.45)
summary.paragraph_format.space_before = Pt(18)
summary.add_run(
    "Structured admissions reference covering programme options, deadlines, personal statement expectations, "
    "A Level, IB and AP requirements, admissions-rate context, and application strategy."
)

doc.add_page_break()

# Table of contents page
toc_title = doc.add_paragraph()
toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
toc_title.paragraph_format.space_before = Pt(54)
toc_title.paragraph_format.space_after = Pt(24)
toc_run = toc_title.add_run("Table of Contents")
toc_run.bold = True
toc_run.font.size = Pt(20)
toc_run.font.name = "Calibri"

for toc_title_text, page in TOC_ENTRIES:
    add_toc_entry(doc, toc_title_text, page)

doc.add_page_break()

# Main content
add_numbered_heading(doc, 1, "Programme Snapshot and Admissions Rates")
add_table(
    doc,
    ["University", "Relevant course/route", "Positioning", "Admissions rate"],
    SNAPSHOT_ROWS,
    [Inches(0.68), Inches(2.15), Inches(2.15), Inches(1.08)],
    font_size=8.3,
    cell_margin=70,
)

doc.add_page_break()

add_numbered_heading(doc, 2, "Course Summaries")
add_table(
    doc,
    ["University", "Course", "Summary"],
    COURSE_SUMMARY_ROWS,
    [Inches(0.75), Inches(2.0), Inches(3.75)],
    font_size=8.6,
    cell_margin=80,
)

doc.add_page_break()

add_numbered_heading(doc, 3, "Application Deadlines")
add_table(
    doc,
    ["University", "Early / priority deadline", "Main / final deadline", "Notes"],
    DEADLINE_ROWS,
    [Inches(0.75), Inches(1.65), Inches(1.85), Inches(2.3)],
    font_size=8.6,
    cell_margin=80,
)

doc.add_page_break()

add_numbered_heading(doc, 4, "Personal Statement Requirements")
add_table(
    doc,
    ["University", "Requirement", "BME-focused writing guidance"],
    PERSONAL_STATEMENT_ROWS,
    [Inches(0.75), Inches(2.75), Inches(3.0)],
    font_size=8.6,
    cell_margin=80,
)

doc.add_page_break()

add_numbered_heading(doc, 5, "A Level Requirements")
add_table(
    doc,
    ["University", "A Level requirement / competitive profile"],
    [[p["uni"], p["alevel"]] for p in PROGRAMMES],
    [Inches(0.85), Inches(5.65)],
)

doc.add_page_break()

add_numbered_heading(doc, 6, "IB Requirements")
add_table(
    doc,
    ["University", "IB requirement / competitive profile"],
    [[p["uni"], p["ib"]] for p in PROGRAMMES],
    [Inches(0.85), Inches(5.65)],
)

doc.add_page_break()

add_numbered_heading(doc, 7, "AP / SAT Requirements")
add_table(
    doc,
    ["University", "AP / SAT requirement / competitive profile"],
    [[p["uni"], p["ap"]] for p in PROGRAMMES],
    [Inches(0.85), Inches(5.65)],
)

doc.add_page_break()

add_numbered_heading(doc, 8, "Planning Takeaways")
for item in [
    "For engineering BME routes, the strongest recurring combination is advanced Mathematics plus at least one laboratory science. Physics is especially useful for engineering-heavy routes; Biology and Chemistry are especially useful for biomedical/health-science routes.",
    "Minimum eligibility is not the same as a realistic offer target. For HKU and HKUST, use the published competitive or mid-50% ranges as better planning signals than the bare minimum.",
    "Where AP is used, AP Calculus plus AP Biology, Chemistry or Physics is the cleanest match to the programme-specific subject expectations.",
    "Recommended strategy: treat HKU as the aspirational first choice if the student is on track for the stated competitive profile; pair it with CUHK or HKUST for a strong research-intensive alternative, and include PolyU and CityU as practical BME options with clear applied-engineering pathways. For AP applicants, prioritise Calculus plus two lab sciences before adding broader STEM electives.",
]:
    doc.add_paragraph(item, style="List Bullet")

for paragraph in doc.paragraphs:
    if paragraph.style.name.startswith("Heading"):
        paragraph.paragraph_format.keep_with_next = True

doc.save(DOCX)
print(DOCX.resolve())
